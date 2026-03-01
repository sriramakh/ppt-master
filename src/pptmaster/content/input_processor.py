"""Unified input processor — auto-detects and extracts text from various formats."""

from __future__ import annotations

from pathlib import Path
from typing import Any


def process_input(source: str | Path) -> str:
    """Auto-detect input type and extract text content.

    Supports: plain text, .pdf, .docx, .csv, .xlsx, .xls, URLs.
    """
    source_str = str(source)

    # URL detection
    if source_str.startswith(("http://", "https://")):
        return _extract_from_url(source_str)

    path = Path(source_str)

    # If it's a file path, process by extension
    if path.exists():
        suffix = path.suffix.lower()
        processors = {
            ".pdf": _extract_from_pdf,
            ".docx": _extract_from_docx,
            ".doc": _extract_from_docx,
            ".csv": _extract_from_csv,
            ".xlsx": _extract_from_excel,
            ".xls": _extract_from_excel,
            ".txt": _extract_from_text,
            ".md": _extract_from_text,
        }
        processor = processors.get(suffix)
        if processor:
            return processor(path)
        # Try reading as text
        return _extract_from_text(path)

    # Otherwise treat as raw text input
    return source_str


def _extract_from_pdf(path: Path) -> str:
    """Extract text from a PDF using pdfplumber."""
    import pdfplumber

    texts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)

            # Also extract tables as formatted text
            tables = page.extract_tables()
            for table in tables:
                if table:
                    texts.append(_format_table(table))

    return "\n\n".join(texts)


def _extract_from_docx(path: Path) -> str:
    """Extract text from a Word document."""
    from docx import Document

    doc = Document(str(path))
    texts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            texts.append(_format_table(rows))

    return "\n\n".join(texts)


def _extract_from_csv(path: Path) -> str:
    """Extract data from a CSV file."""
    import pandas as pd

    df = pd.read_csv(path)
    return _dataframe_to_text(df)


def _extract_from_excel(path: Path) -> str:
    """Extract data from an Excel file."""
    import pandas as pd

    sheets = pd.read_excel(path, sheet_name=None)
    texts: list[str] = []
    for sheet_name, df in sheets.items():
        texts.append(f"## Sheet: {sheet_name}")
        texts.append(_dataframe_to_text(df))
    return "\n\n".join(texts)


def _extract_from_text(path: Path) -> str:
    """Read a plain text file."""
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_from_url(url: str) -> str:
    """Extract main content from a web page using trafilatura."""
    import trafilatura

    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        raise ValueError(f"Could not fetch URL: {url}")

    text = trafilatura.extract(downloaded, include_tables=True, include_links=False)
    if not text:
        raise ValueError(f"Could not extract content from URL: {url}")

    return text


def _dataframe_to_text(df: Any) -> str:
    """Convert a pandas DataFrame to a descriptive text summary."""
    lines: list[str] = []

    lines.append(f"Data with {len(df)} rows and {len(df.columns)} columns.")
    lines.append(f"Columns: {', '.join(str(c) for c in df.columns)}")

    # Summary statistics for numeric columns
    numeric_cols = df.select_dtypes(include=["number"]).columns
    if len(numeric_cols) > 0:
        lines.append("\nNumeric summary:")
        lines.append(df[numeric_cols].describe().to_string())

    # First few rows as sample
    lines.append("\nSample data:")
    lines.append(df.head(10).to_string())

    return "\n".join(lines)


def _format_table(rows: list[list[Any]]) -> str:
    """Format a table as text."""
    if not rows:
        return ""
    lines: list[str] = []
    for row in rows:
        if row:
            lines.append(" | ".join(str(cell) if cell else "" for cell in row))
    return "\n".join(lines)
